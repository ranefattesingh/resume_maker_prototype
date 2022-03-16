import docx
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.table import _Cell
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import json
import os


def set_cell_margins(cell: _Cell, **kwargs):
    """
    cell:  actual cell instance you want to modify
    usage:
        set_cell_margins(cell, top=50, start=50, bottom=50, end=50)
    provided values are in twentieths of a point (1/1440 of an inch).
    read more here: http://officeopenxml.com/WPtableCellMargins.php
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')

    for m in [
        "top",
        "start",
        "bottom",
        "end",
    ]:
        if m in kwargs:
            node = OxmlElement("w:{}".format(m))
            node.set(qn('w:w'), str(kwargs.get(m)))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)

    tcPr.append(tcMar)


def set_cell_border(cell: _Cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

f = open('profile.json')
profile = json.load(f)

name = profile['name']
bio_data = profile['bio_data']
data = profile['data']

bold_label = "bold_value"

def make_table_row_bold(cell):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size= Pt(12)

def change_table_row_color(cell):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(0x69, 0x69, 0x69)

def set_cell_width(cell, width):
    cell.width = Inches(width)

def set_row_height(row):
    row.height = Cm(0.5)
    
def hide_cell_border(cell):
    set_cell_border(
        cell,
        top={"sz": 1, "val": "single", "color": "#FFFFFF"},
        bottom={"sz": 1, "val": "single", "color": "#FFFFFF"},
        start={"sz": 1, "val": "single", "color": "#FFFFFF"},
        end={"sz": 1, "val": "single", "color": "#FFFFFF"},
    )

def hide_table_border(table):
    for row in table.rows:
        for cell in row.cells:
            hide_cell_border(cell)

doc = docx.Document()
default_styles = doc.styles
default_font = "Calibri (Body)"
doc.styles["Normal"].font.name = default_font

heading = doc.add_table(rows = 0, cols = 2)
heading.style = 'Table Grid'
row_cells = heading.add_row().cells
set_cell_width(row_cells[0], 3)
set_cell_width(row_cells[1], 6)
name = name.upper().split(' ')
name_font_size = 21
first_name_part1 = row_cells[0].paragraphs[0].add_run(name[0][0])
first_name_part1.font.size = Pt(name_font_size)
first_name_part2 = row_cells[0].paragraphs[0].add_run(name[0][1:])
first_name_part2.font.size = Pt(name_font_size - 2)
row_cells[0].paragraphs[0].add_run(" ")
last_name_part1 = row_cells[0].paragraphs[0].add_run(name[1][0])
last_name_part1.font.size = Pt(name_font_size)
last_name_part2 = row_cells[0].paragraphs[0].add_run(name[1][1:])
last_name_part2.font.size = Pt(name_font_size - 2)

row_cells[1]._element.clear_content()
bios = row_cells[1].add_table(rows = 0, cols = 2)
set_cell_margins(row_cells[1], top=0, start=0, bottom=0, end=0)
bios.style = 'Table Grid'
for bio in bio_data:
    bio_cells = bios.add_row().cells
    set_cell_width(bio_cells[0], 0.3)
    set_cell_width(bio_cells[1], 6)
    bio_paragraph = bio_cells[0].paragraphs[0]
    try:
        bio_paragraph_run = bio_paragraph.add_run()
        bio_paragraph_run.add_picture(f'./icons/{bio}.png', width=Cm(0.5), height=Cm(0.5))
    except FileNotFoundError:
        print(f'{bio} Icon is not defined')
    bio_cells[1].text = bio_data[bio]
    hide_cell_border(bio_cells[0])
    hide_cell_border(bio_cells[1])

hide_cell_border(row_cells[0])
hide_cell_border(row_cells[1])

space = doc.add_paragraph()
space.paragraph_format.line_spacing = Cm(0.1)
gap_table = doc.add_table(rows = 0, cols = 1)
gap_table.style = "Table Grid"
gap_table_cell = gap_table.add_row().cells[0]
set_cell_border(
    gap_table_cell,
    top={"sz": 12, "val": "single", "shadow": "true"},
    bottom={"sz": 1, "val": "single", "color": "#FFFFFF"},
    start={"sz": 1, "val": "single", "color": "#FFFFFF"},
    end={"sz": 1, "val": "single", "color": "#FFFFFF"},
)

#MERGE ROWS
for i in range(0, len(heading.rows) - 1):
    row1 = heading.cell(i, 0)
    row2 = heading.cell(i + 1, 0)
    row1.merge(row2)

#CREATE RESUME ROWS
level0 = doc.add_table(rows = 0, cols = 2)
level0.style = 'Table Grid'
for item in data:
    #THIS LOOP CREATE OUTERMOST CONTAINER AND FILLS DATA IN LEFT
    group = data[item]
    level0_row = level0.add_row()
    level0_row_cells = level0_row.cells
    level0_row_cells[0].text = item.title()
    change_table_row_color(level0_row_cells[0])
    make_table_row_bold(level0_row_cells[0])
    set_cell_width(level0_row_cells[0], 3.44)
    set_cell_width(level0_row_cells[1], 12.18)
    hide_cell_border(level0_row_cells[1])
    set_cell_border(
        level0_row_cells[0],
        top={"sz": 1, "val": "single", "color": "#FFFFFF"},
        bottom={"sz": 1, "val": "single", "color": "#FFFFFF"},
        start={"sz": 1, "val": "single", "color": "#FFFFFF"},
        end={"sz": 5, "val": "single", "shadow": "true"},
    )
    level0_row_cells[1]._element.clear_content()
    level1 = level0_row_cells[1].add_table(rows = 0, cols = 1)
    level1.style = 'Table Grid'
    for item1 in group:
        #THIS LOOP CREATES INNER CONTAINER FOR DATA IN RIGHT DIRECTION
        group_item = group[item1]
        level1_row = level1.add_row()
        level1_row_cells = level1_row.cells
        level1_row_cells[0]._element.clear_content()
        hide_cell_border(level1_row_cells[0])
        level2 = level1_row_cells[0].add_table(rows = 0, cols = 1)
        set_cell_margins(level1_row_cells[0], top=0, start=0, bottom=0, end=0)
        level2.style = 'Table Grid'
        

        if isinstance(group_item, str):
            entry = group_item
            level2_row = level2.add_row()
            level2_row_cells = level2_row.cells
            set_cell_margins(level2_row_cells[0], top=0, start=0, bottom=0,end=0)
            level2_row_cells[0].text = entry
            change_table_row_color(level2_row_cells[0])
            hide_cell_border(level2_row_cells[0])

            if list(group.keys())[-1]==item1:
                space_row_cells = level1.add_row().cells
                hide_cell_border(space_row_cells[0])
        else:     
            for item2 in group_item:
                #THIS LOOP FILLS DATA IN RIGHT CONTAINER
                entry = group_item[item2]
                level2_row = level2.add_row()
                level2_row_cells = level2_row.cells
                set_cell_margins(level2_row_cells[0], top=0, start=0, bottom=0, end=0)
                change_table_row_color(level2_row_cells[0])
                level2_row_cells[0]._element.clear_content()
                if bold_label in item2:
                    level2_row_cells[0].text = entry
                    make_table_row_bold(level2_row_cells[0])
                else:
                    level2_row_cells[0].text = entry
                    change_table_row_color(level2_row_cells[0])
                hide_cell_border(level2_row_cells[0])

        if not isinstance(group_item, str):
            space_row_cells = level1.add_row().cells
        hide_cell_border(space_row_cells[0])
doc.save("resume.docx")
print("Resume is saved in project directory", os.getcwd())